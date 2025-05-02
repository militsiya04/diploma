import base64
import hashlib
import io
import os
import random
import re
import subprocess
import sys
from datetime import datetime, date, timedelta
from io import BytesIO

import cv2
import numpy as np
import pandas as pd
from deepface import DeepFace
from docx import Document
from flask import (
    Flask,
    flash,
    jsonify,
    redirect,
    render_template,
    request,
    send_file,
    send_from_directory,
    session,
    url_for,
)
from PIL import Image
from reportlab.lib.pagesizes import letter
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas
from services.admin_setup import (
    check_and_generate_admin_link,
    generate_registration_link,
)
from services.captcha_generator import generate_captcha_image, generate_captcha_text
from services.document_reader import extract_text_from_docx, extract_text_from_pdf
from services.email_otp_service import configure_mail, send_otp_email
from services.init_database import get_db_connection, init_database
from services.otp_checking import store_otp, validate_otp
from services.send_sms import send_sms
from utils import (
    allowed_file,
    check_password,
    generate_session_token,
    hash_password,
    is_fully_authenticated,
    login_required_with_timeout,
    roles_required,
)
from werkzeug.utils import secure_filename

from services.crypto import (
    decrypt_rsa,
    encrypt_rsa,
    load_private_key,
    load_public_key,
)
from services.init_rsa_keys import generate_rsa_keys

app = Flask(__name__)
app.secret_key = "d9f9a8b7e5a4422aa1c8cf59d6d22e80"

UPLOAD_FOLDER = "uploads"
DATABASE_FOLDER = "database"
SERVER_DATABASE_FOLDER = "server_database"
EXCEL_FOLDER = os.path.join(SERVER_DATABASE_FOLDER, "excel_files")
VERIFICATION_PHOTOS_FOLDER = os.path.join(SERVER_DATABASE_FOLDER, "verification_photos")

app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(DATABASE_FOLDER, exist_ok=True)
os.makedirs(SERVER_DATABASE_FOLDER, exist_ok=True)
os.makedirs(EXCEL_FOLDER, exist_ok=True)
os.makedirs(VERIFICATION_PHOTOS_FOLDER, exist_ok=True)

configure_mail(app)
init_database()
check_and_generate_admin_link(get_db_connection())
generate_rsa_keys()


# ----- AUTHENTICATION CYCLE START -----
@app.route("/captcha")
def captcha():
    captcha_text = generate_captcha_text()
    session["captcha"] = captcha_text
    img_io = generate_captcha_image(captcha_text)
    return send_file(img_io, mimetype="image/png")


@app.route("/register/<token>", methods=["GET", "POST"])
def register_user(token: str):
    if is_fully_authenticated():
        return redirect(url_for("redirect_user"))

    conn = get_db_connection()
    cursor = conn.cursor()

    cursor.execute(
        "SELECT token, expiry, used, role FROM registration_tokens WHERE token = ?",
        (token,),
    )
    token_data = cursor.fetchone()

    if not token_data:
        return "⛔ Посилання недійсне або не існує.", 403
    if token_data[2]:
        return "⛔ Це посилання вже використано.", 403
    if datetime.now() > token_data[1]:
        return "⛔ Посилання протерміноване.", 403

    role = token_data[3]

    if role == "admin":
        cursor.execute("SELECT id FROM users WHERE position = 'admin'")
        if cursor.fetchone():
            conn.close()
            return (
                "⚠️ Адміністратор уже існує в системі. Повторна реєстрація неможлива.",
                403,
            )

    if request.method == "POST":
        login = request.form["login"]
        password = request.form["password"]
        email = request.form["email"]
        phone = request.form["phone"]
        first_name = request.form["first_name"]
        surname = request.form["surname"]
        user_captcha = request.form["captcha"].strip().lower()
        session_captcha = session.get("captcha", "").strip().lower()

        if user_captcha != session_captcha:
            flash(" Невірна капча!", "error")
            return redirect(request.url)

        session.pop("captcha", None)

        email_hash = hashlib.sha256(email.encode()).hexdigest()
        phone_hash = hashlib.sha256(phone.encode()).hexdigest()

        cursor.execute(
            "SELECT * FROM users WHERE login = ? OR email_hash = ? OR phone_hash = ?",
            (login, email_hash, phone_hash),
        )
        if cursor.fetchone():
            conn.close()
            flash(
                "❗ Користувач з таким логіном, email або телефоном уже існує!", "error"
            )
            return redirect(request.url)

        public_key = load_public_key()
        encrypted_email = encrypt_rsa(email, public_key)
        encrypted_phone = encrypt_rsa(phone, public_key)
        encrypted_first_name = encrypt_rsa(first_name, public_key)
        encrypted_surname = encrypt_rsa(surname, public_key)

        hashed_password = hash_password(password)

        cursor.execute(
            """
            INSERT INTO users (position, login, password, email, phone, first_name, surname, email_hash, phone_hash)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                role,
                login,
                hashed_password,
                encrypted_email,
                encrypted_phone,
                encrypted_first_name,
                encrypted_surname,
                email_hash,
                phone_hash,
            ),
        )

        cursor.execute("SELECT @@IDENTITY")
        user_id = cursor.fetchone()[0]

        photo = request.files.get("photo")
        if photo and photo.filename.lower().endswith((".jpg", ".jpeg")):
            img = Image.open(photo)
            save_path = f"server_database/verification_photos/{user_id}.jpg"
            os.makedirs(os.path.dirname(save_path), exist_ok=True)
            img.convert("RGB").save(save_path, "JPEG")
        else:
            flash(" Фото має бути у форматі JPG!", "error")
            conn.rollback()
            return redirect(request.url)

        cursor.execute(
            "UPDATE registration_tokens SET used = 1 WHERE token = ?", (token,)
        )
        conn.commit()
        conn.close()

        flash(f"Користувач з роллю '{role}' успішно зареєстрований!", "success")
        return redirect(url_for("login"))

    return render_template("register.html", token=token, role=role)


@app.route("/", methods=["GET", "POST"])
def login():
    if is_fully_authenticated():
        return redirect(url_for("redirect_user"))

    if request.method == "POST":
        login = request.form["login"]
        password = request.form["password"]
        user_captcha = request.form["captcha"].strip().lower()
        session_captcha = session.get("captcha", "").strip().lower()

        if user_captcha != session_captcha:
            flash("Помилка: Невірна капча!", "error")
            return redirect(url_for("login"))

        session.pop("captcha", None)

        conn = get_db_connection()
        cursor = conn.cursor()

        cursor.execute(
            "SELECT id, position, email, phone, password FROM users WHERE login=?",
            (login,),
        )
        user = cursor.fetchone()
        conn.close()

        if user and check_password(password, user[4]):
            private_key = load_private_key()

            session["user_id"] = user[0]
            session["user_position"] = user[1]
            session["user_email"] = decrypt_rsa(user[2], private_key)
            session["user_phone"] = decrypt_rsa(user[3], private_key)
            return redirect(url_for("auth_options"))
        else:
            flash("Помилка: Неправильні дані!", "error")

    return render_template("login.html")


@app.route("/auth-options", methods=["GET", "POST"])
def auth_options():
    if is_fully_authenticated():
        return redirect(url_for("redirect_user"))
    if "user_id" not in session:
        return redirect(url_for("login"))

    return render_template("auth_options.html")


@app.route("/authenticate-email", methods=["POST"])
def authenticate_email():
    if "user_id" not in session:
        return redirect(url_for("login"))
    if "user_email" not in session:
        flash("Електронної пошти не знайдено.", "error")
        return redirect(url_for("auth_options"))

    email = session["user_email"]
    user_id = session["user_id"]
    code = "{:06d}".format(random.randint(0, 999999))

    conn = get_db_connection()
    cursor = conn.cursor()
    store_otp(cursor, user_id, code, "email")
    conn.commit()
    conn.close()

    send_otp_email(email, code)

    return redirect(url_for("verify_email"))


@app.route("/verify-email", methods=["GET", "POST"])
def verify_email():
    if is_fully_authenticated():
        return redirect(url_for("redirect_user"))
    if "user_id" not in session:
        return redirect(url_for("login"))

    conn = get_db_connection()
    cursor = conn.cursor()

    if request.method == "POST":
        entered_otp = request.form["otp"]
        if validate_otp(cursor, session["user_id"], entered_otp, "email"):
            conn.commit()
            conn.close()
            session["session_token"] = generate_session_token(
                session["user_id"], session["user_position"]
            )
            return redirect(url_for("redirect_user"))
        else:
            flash("Недійсний або прострочений OTP. Спробуйте ще раз.", "error")

    conn.close()
    return render_template("verify_email.html")


@app.route("/authenticate-phone", methods=["POST"])
def authenticate_phone():
    if "user_id" not in session:
        return redirect(url_for("login"))
    if "user_phone" not in session or not session["user_phone"]:
        flash("Помилка: Телефон не знайдено.", "error")
        return redirect(url_for("auth_options"))
    return redirect(url_for("verify_phone"))


@app.route("/verify_phone", methods=["GET", "POST"])
def verify_phone():
    if "user_id" not in session:
        return redirect(url_for("login"))
    user_id = session["user_id"]

    conn = get_db_connection()
    cursor = conn.cursor()

    if request.method == "POST":
        entered_otp = request.form["otp"]
        if validate_otp(cursor, user_id, entered_otp, "phone"):
            conn.commit()
            conn.close()
            session["session_token"] = generate_session_token(
                user_id, session["user_position"]
            )
            return redirect(url_for("redirect_user"))
        else:
            flash("Недійсний або прострочений OTP.", "error")
    else:
        code = "{:06d}".format(random.randint(0, 999999))
        store_otp(cursor, user_id, code, "phone")
        phone_number = session.get("user_phone")
        if phone_number:
            result = send_sms([phone_number], f"Ваш код подтверждения: {code}")
            if "error" in result:
                flash("Помилка при надсиланні SMS: " + result["error"], "error")
            else:
                flash("Код відправлено на номер телефону.", "info")
        else:
            flash("Номер телефону не знайдено у сесії.", "error")
        conn.commit()

    conn.close()
    return render_template("verify_phone.html")


@app.route("/verify_face", methods=["GET", "POST"])
def verify_face():
    if request.method == "POST":
        user_id = session.get("user_id")
        if not user_id:
            flash("Ідентифікатор користувача не знайдено в сеансі", "error")
            return redirect(url_for("login"))

        photo_data = request.form.get("photo")
        if not photo_data:
            flash("Фото не отримано", "error")
            return redirect(url_for("verify_face"))

        img_str = re.search(r"base64,(.*)", photo_data).group(1)
        img_bytes = base64.b64decode(img_str)
        np_arr = np.frombuffer(img_bytes, np.uint8)
        frame = cv2.imdecode(np_arr, cv2.IMREAD_COLOR)

        input_photo_path = f"server_database/verification_photos/temp_{user_id}.jpg"
        reference_photo_path = f"server_database/verification_photos/{user_id}.jpg"
        cv2.imwrite(input_photo_path, frame)

        if not os.path.exists(reference_photo_path):
            flash("Фото для перевірки не знайдено!", "error")
            os.remove(input_photo_path)
            return redirect(url_for("verify_face"))

        try:
            result = DeepFace.verify(
                img1_path=input_photo_path,
                img2_path=reference_photo_path,
                enforce_detection=True,
            )
            os.remove(input_photo_path)

            if result["verified"]:
                session["session_token"] = generate_session_token(
                    session["user_id"], session["user_position"]
                )
                return redirect(url_for("redirect_user"))
            else:
                flash(" Особа не збігається.", "error")
                return redirect(url_for("verify_face"))

        except Exception as e:
            flash(f"Помилка порівняння", "error")
            if os.path.exists(input_photo_path):
                os.remove(input_photo_path)
            return redirect(url_for("verify_face"))

    return render_template("verify_face.html")


@app.route("/redirect-user")
def redirect_user():
    if not is_fully_authenticated():
        return redirect(url_for("login"))

    if session["user_position"] == "patient":
        return redirect(url_for("dashboard"))
    elif session["user_position"] in ["doctor", "admin"]:
        return redirect(url_for("meddashboard"))

    flash("Помилка: Невідома роль користувача.", "error")
    return redirect(url_for("login"))


@app.route("/logout")
def logout():
    session.clear()
    flash("Ви вийшли з системи.", "info")
    return redirect(url_for("login"))


# ----- AUTHENTICATION CYCLE END -----


@app.route("/forgot-password", methods=["GET", "POST"])
def forgot_password():
    if request.method == "POST":
        login = request.form["login"]

        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute("SELECT id FROM users WHERE login = ?", (login,))
        user = cursor.fetchone()
        conn.close()

        if user:
            user_id = user[0]
            session["reset_user_id"] = user_id
            return redirect(url_for("verify_reset_face"))
        else:
            flash("Користувача з таким логіном не знайдено.", "error")

    return render_template("forgot_password.html")


@app.route("/verify-reset-face", methods=["GET", "POST"])
def verify_reset_face():
    if "reset_user_id" not in session:
        return redirect(url_for("forgot_password"))

    user_id = session["reset_user_id"]

    if request.method == "POST":
        photo_data = request.form.get("photo")
        if not photo_data:
            flash("Фото не отримано", "error")
            return redirect(url_for("verify_reset_face"))

        img_str = re.search(r"base64,(.*)", photo_data).group(1)
        img_bytes = base64.b64decode(img_str)
        np_arr = np.frombuffer(img_bytes, np.uint8)
        frame = cv2.imdecode(np_arr, cv2.IMREAD_COLOR)

        input_photo_path = (
            f"server_database/verification_photos/temp_reset_{user_id}.jpg"
        )
        reference_photo_path = f"server_database/verification_photos/{user_id}.jpg"
        cv2.imwrite(input_photo_path, frame)

        if not os.path.exists(reference_photo_path):
            flash("Фото користувача не знайдено!", "error")
            os.remove(input_photo_path)
            return redirect(url_for("verify_reset_face"))

        try:
            result = DeepFace.verify(
                img1_path=input_photo_path,
                img2_path=reference_photo_path,
                enforce_detection=True,
            )
            os.remove(input_photo_path)

            if result["verified"]:
                # Генерируем токен для сброса пароля
                reset_token = secrets.token_urlsafe(32)
                session["reset_token"] = reset_token

                # Редиректим на страницу сброса пароля с токеном в URL
                return redirect(url_for("reset_password", token=reset_token))
            else:
                flash("Обличчя не співпадає.", "error")
                return redirect(url_for("verify_reset_face"))

        except Exception as e:
            flash(f"Помилка перевірки обличчя: {str(e)}", "error")
            if os.path.exists(input_photo_path):
                os.remove(input_photo_path)
            return redirect(url_for("verify_reset_face"))

    return render_template("verify_reset_face.html")


@app.route("/reset-password", methods=["GET", "POST"])
def reset_password():
    token = request.args.get("token")

    # Проверяем токен
    if not token or session.get("reset_token") != token:
        flash("Доступ заборонено.", "error")
        return redirect(url_for("forgot_password"))

    if request.method == "POST":
        new_password = request.form["password"]
        hashed_password = hash_password(new_password)

        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute(
            "UPDATE users SET password = ? WHERE id = ?",
            (hashed_password, session["reset_user_id"]),
        )
        conn.commit()
        conn.close()

        # Очищаем всё после смены пароля
        session.pop("reset_user_id", None)
        session.pop("reset_token", None)

        flash("Пароль успішно змінено! Увійдіть з новим паролем.", "success")
        return redirect(url_for("login"))

    return render_template("reset_password.html")


@app.route("/generate-links", methods=["GET", "POST"])
@login_required_with_timeout()
@roles_required("admin", "doctor")
def generate_links():
    conn = get_db_connection()
    cursor = conn.cursor()

    current_user_role = session["user_position"]
    allowed_roles = []

    if current_user_role == "admin":
        allowed_roles = ["doctor", "patient"]
    elif current_user_role == "doctor":
        allowed_roles = ["patient"]

    if request.method == "POST":
        selected_role = request.form.get("role")
        if selected_role not in allowed_roles:
            flash(" Ви не маєте прав створювати користувачів з цією роллю!", "error")
            return redirect(url_for("generate_links"))

        generate_registration_link(conn, selected_role, hours_valid=24)
        flash(f"Посилання для {selected_role} згенеровано!", "success")
        return redirect(url_for("generate_links"))

    cursor.execute(
        """
        SELECT token, role, expiry
        FROM registration_tokens
        WHERE used = 0 AND expiry > ?
        ORDER BY expiry ASC
    """,
        (datetime.now().strftime("%Y-%m-%d %H:%M:%S"),),
    )

    token_data = cursor.fetchall()
    conn.close()

    base_url = request.host_url.rstrip("/")
    tokens = [
        {
            "url": f"{base_url}/register/{row[0]}",
            "role": row[1],
            "expiry": row[2],
        }
        for row in token_data
    ]

    return render_template(
        "generate_links.html", allowed_roles=allowed_roles, tokens=tokens
    )


@app.route("/add_pulse/<int:patient_id>", methods=["POST"])
@login_required_with_timeout()
@roles_required("admin", "doctor")
def add_pulse(patient_id):
    try:
        pulse = request.form["pulse"]
        selected_date_str = request.form["selected_date"]
        selected_date = datetime.strptime(selected_date_str, "%Y-%m-%d")

        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute(
            "INSERT INTO pulse (user_id, pulse, date_when_created) VALUES (?, ?, ?)",
            (patient_id, pulse, selected_date),
        )
        conn.commit()
        conn.close()
        flash("Пульс додано успішно!")
    except Exception as e:
        flash(f"Помилка бази даних: {e}", "error")

    return redirect(url_for("patient_dashboard", patient_id=patient_id))


@app.route("/add_dispersion/<int:patient_id>", methods=["POST"])
@login_required_with_timeout()
@roles_required("admin", "doctor")
def add_dispersion(patient_id):
    try:
        pulse = request.form["pulse"]
        pressure = request.form["pressure"]
        oxygen_level = request.form["oxygen_level"]
        weight = request.form["weight"]
        sugar = request.form["sugar"]
        temperature = request.form["temperature"]
        selected_date_str = request.form["selected_date"]

        selected_date = datetime.strptime(selected_date_str, "%Y-%m-%d")

        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute(
            """
            INSERT INTO dispersion (user_id, pulse, pressure, oxygen_level, weight, sugar, temperature, date_when_created)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?)
        """,
            (
                patient_id,
                pulse,
                pressure,
                oxygen_level,
                weight,
                sugar,
                temperature,
                selected_date,
            ),
        )
        conn.commit()
        conn.close()

        flash("Дані успішно збережені!")
    except Exception as e:
        flash(f"Помилка бази даних: {e}", "error")

    return redirect(url_for("patient_dashboard", patient_id=patient_id))


@app.route("/add_was/<int:patient_id>", methods=["POST"])
@login_required_with_timeout()
@roles_required("admin", "doctor")
def add_was(patient_id):
    try:
        weight = request.form["weight"]
        sugar = request.form["sugar"]
        selected_date_str = request.form["selected_date"]
        selected_date = datetime.strptime(selected_date_str, "%Y-%m-%d")

        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute(
            "INSERT INTO WaS (user_id, weight, sugar, date_when_created) VALUES (?, ?, ?, ?)",
            (patient_id, weight, sugar, selected_date),
        )
        conn.commit()
        conn.close()
        flash("Дані про вагу та цукор додано успішно!")
    except Exception as e:
        flash(f"Помилка бази даних: {e}", "error")

    return redirect(url_for("patient_dashboard", patient_id=patient_id))


@app.route("/add_pressure/<int:patient_id>", methods=["POST"])
@login_required_with_timeout()
@roles_required("admin", "doctor")
def add_pressure(patient_id):
    try:
        bpressure = request.form["bpressure"]
        apressure = request.form["apressure"]
        selected_date_str = request.form["selected_date"]
        selected_date = datetime.strptime(selected_date_str, "%Y-%m-%d")

        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute(
            "INSERT INTO pressure (user_id, bpressure, apressure, date_when_created) VALUES (?, ?, ?, ?)",
            (patient_id, bpressure, apressure, selected_date),
        )
        conn.commit()
        conn.close()
        flash("Тиск додано успішно!")
    except Exception as e:
        flash(f"Помилка бази даних: {e}", "error")

    return redirect(url_for("patient_dashboard", patient_id=patient_id))


@app.route("/dashboard")
@login_required_with_timeout()
@roles_required("patient")
def dashboard():
    user_id = session["user_id"]

    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT position FROM users WHERE id = ?", (user_id,))
    position = cursor.fetchone()

    if not position:
        conn.close()
        flash("Користувача не знайдено.", "error")
        return redirect(url_for("login"))

    cursor.execute(
        "SELECT first_name, surname, phone, info FROM users WHERE id = ?",
        (user_id,),
    )
    user = cursor.fetchone()
    conn.close()

    private_key = load_private_key()
    first_name = decrypt_rsa(user[0], private_key)
    surname = decrypt_rsa(user[1], private_key)
    phone = decrypt_rsa(user[2], private_key)
    info = decrypt_rsa(user[3], private_key) if user[3] else "N/A"

    patient_folder = os.path.join("server_database/excel_files/", str(user_id))
    files = os.listdir(patient_folder) if os.path.exists(patient_folder) else []

    return render_template(
        "dashboard.html",
        user={
            "patient_id": user_id,
            "first_name": first_name,
            "surname": surname,
            "phone": phone,
            "info": info,
        },
        files=files,
    )


@app.route("/download-info/<string:format>/<int:patient_id>")
@login_required_with_timeout()
@roles_required("admin", "doctor", "patient")
def download_info(format, patient_id):
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute(
        "SELECT first_name, surname, email, phone, position, info FROM users WHERE id = ?",
        (patient_id,),
    )
    user = cursor.fetchone()
    if not user:
        conn.close()
        flash("Інформація про користувача не знайдена.", "error")
        return redirect(url_for("dashboard"))

    private_key = load_private_key()
    first_name = decrypt_rsa(user[0], private_key)
    surname = decrypt_rsa(user[1], private_key)
    email = decrypt_rsa(user[2], private_key)
    phone = decrypt_rsa(user[3], private_key)
    position = user[4]
    info = decrypt_rsa(user[5], private_key) if user[5] else "N/A"
    conn.close()

    # --- Медичні дані ---
    conn_med = get_db_connection()
    pulse_df = pd.read_sql(
        f"SELECT pulse FROM pulse WHERE user_id = {patient_id}", conn_med
    )
    pressure_df = pd.read_sql(
        f"SELECT bpressure, apressure FROM Pressure WHERE user_id = {patient_id}",
        conn_med,
    )
    weight_sugar_df = pd.read_sql(
        f"SELECT weight, sugar FROM WaS WHERE user_id = {patient_id}", conn_med
    )
    conn_med.close()

    # Аналіз пульсу
    pulse_text = ""
    if not pulse_df.empty:
        pulse_avg = pulse_df["pulse"].dropna().mean()
        pulse_text = f"Середній пульс: {pulse_avg:.2f} уд/хв"
        if pulse_avg < 60 or pulse_avg > 100:
            pulse_text += "\nПульс виходить за межі норми (60-100 уд/хв)!"

    # Аналіз тиску
    pressure_text = ""
    if not pressure_df.empty:
        bp_avg = pressure_df["bpressure"].dropna().mean()
        ap_avg = pressure_df["apressure"].dropna().mean()
        pressure_text = (
            f"Середній початковий тиск: {bp_avg:.2f} мм рт.ст.\n"
            f"Середній після лікування: {ap_avg:.2f} мм рт.ст."
        )
        if bp_avg < 90 or bp_avg > 140:
            pressure_text += "\nПочатковий тиск виходить за межі норми (90–140)!"
        if ap_avg < 90 or ap_avg > 140:
            pressure_text += "\nТиск після лікування виходить за межі норми (90–140)!"

    # Оцінка ефективності лікування
    treatment_effect_text = ""
    if not pressure_df.empty:
        before = pressure_df["bpressure"].dropna()
        after = pressure_df["apressure"].dropna()

        if len(before) >= 2 and len(before) == len(after):
            avg_before = before.mean()
            avg_after = after.mean()
            delta = avg_before - avg_after

            treatment_effect_text = (
                f"Ефективність лікування (аналіз тиску):\n"
                f"  - Середній тиск до лікування: {avg_before:.2f} мм рт.ст.\n"
                f"  - Середній тиск після лікування: {avg_after:.2f} мм рт.ст.\n"
                f"  - Різниця: {delta:.2f} мм рт.ст.\n"
            )

            if delta > 0:
                treatment_effect_text += (
                    "Спостерігається зниження тиску після лікування.\n"
                )
            elif delta < 0:
                treatment_effect_text += "Після лікування тиск підвищився.\n"
            else:
                treatment_effect_text += "Зміни тиску не виявлено.\n"
        else:
            treatment_effect_text = (
                "Ефективність лікування: недостатньо парних даних для аналізу.\n"
            )

    # Аналіз ваги та цукру
    weight_sugar_text = ""
    if not weight_sugar_df.empty:
        weight_sugar_df.columns = [
            col.strip().lower() for col in weight_sugar_df.columns
        ]
        weight_sugar_df["parsed_sugar"] = (
            weight_sugar_df["sugar"]
            .astype(str)
            .str.replace(",", ".", regex=False)
            .astype(float)
        )
        avg_weight = weight_sugar_df["weight"].dropna().mean()
        avg_sugar = weight_sugar_df["parsed_sugar"].dropna().mean()

        weight_sugar_text = (
            f"Середня вага: {avg_weight:.2f} кг\n"
            f"Середній рівень цукру: {avg_sugar:.2f} ммоль/л"
        )

        if avg_sugar < 3.9 or avg_sugar > 7.8:
            weight_sugar_text += (
                "\nРівень цукру виходить за межі норми (3.9–7.8 ммоль/л)!"
            )

    # Дисперсія за останні 30 днів
    dispersion_text = ""
    try:
        conn_disp = get_db_connection()
        end_date = datetime.now()
        start_date = end_date - timedelta(days=30)
        query = (
            f"SELECT pulse, pressure, oxygen_level, weight, sugar, temperature "
            f"FROM dispersion WHERE user_id = {patient_id} "
            f"AND date_when_created BETWEEN ? AND ?"
        )
        df_disp = pd.read_sql(query, conn_disp, params=[start_date, end_date])
        conn_disp.close()

        if not df_disp.empty:
            df_disp["sugar"] = pd.to_numeric(df_disp["sugar"], errors="coerce")
            df_disp["temperature"] = pd.to_numeric(
                df_disp["temperature"], errors="coerce"
            )

            dispersion_text = "Дисперсія за останні 30 днів:\n"
            for column in [
                "pulse",
                "pressure",
                "oxygen_level",
                "weight",
                "sugar",
                "temperature",
            ]:
                values = df_disp[column].dropna()
                if not values.empty:
                    dispersion_value = values.var()
                    dispersion_text += (
                        f"{column.capitalize()}: {dispersion_value:.2f}\n"
                    )
                else:
                    dispersion_text += f"{column.capitalize()}: Немає даних\n"
        else:
            dispersion_text = "Дисперсія за останні 30 днів: Немає даних\n"

    except Exception as e:
        dispersion_text = f"Помилка при обчисленні дисперсії: {e}"

    # --- Формування звіту ---
    if format == "pdf":
        base_dir = os.path.dirname(os.path.abspath(__file__))
        font_path = os.path.join(base_dir, "static", "fonts", "free-sans.ttf")
        pdfmetrics.registerFont(TTFont("FreeSans", font_path))

        pdf_file = BytesIO()
        c = canvas.Canvas(pdf_file, pagesize=letter)
        c.setFont("FreeSans", 12)

        # Текст
        c.drawString(100, 750, "Інформація про користувача:")
        c.drawString(100, 730, f"Ім'я: {first_name}")
        c.drawString(100, 710, f"Прізвище: {surname}")
        c.drawString(100, 690, f"Email: {email}")
        c.drawString(100, 670, f"Телефон: {phone}")
        c.drawString(100, 650, f"Посада: {position}")
        c.drawString(100, 630, f"Інформація: {info}")

        # Медичні дані
        y = 600
        for line in (
            pulse_text,
            pressure_text,
            treatment_effect_text,
            weight_sugar_text,
        ):
            for subline in line.split("\n"):
                c.drawString(100, y, subline)
                y -= 20
            y -= 10

        for subline in dispersion_text.split("\n"):
            c.drawString(100, y, subline)
            y -= 20
        y -= 10

        c.save()
        pdf_file.seek(0)

        return send_file(
            pdf_file,
            as_attachment=True,
            download_name="user_info.pdf",
            mimetype="application/pdf",
        )

    elif format == "docx":
        doc = Document()
        doc.add_heading("Інформація про користувача", level=1)
        doc.add_paragraph(f"Ім'я: {first_name}")
        doc.add_paragraph(f"Прізвище: {surname}")
        doc.add_paragraph(f"Email: {email}")
        doc.add_paragraph(f"Телефон: {phone}")
        doc.add_paragraph(f"Посада: {position}")
        doc.add_paragraph(f"Інформація: {info}")

        doc.add_heading("Медичні дані", level=2)
        doc.add_paragraph(pulse_text)
        doc.add_paragraph(pressure_text)
        doc.add_paragraph(treatment_effect_text)
        doc.add_paragraph(weight_sugar_text)
        doc.add_paragraph(dispersion_text)

        doc_file = BytesIO()
        doc.save(doc_file)
        doc_file.seek(0)

        return send_file(
            doc_file,
            as_attachment=True,
            download_name="user_info.docx",
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    flash("Непідтримуваний формат файлу.", "error")
    return redirect(url_for("dashboard"))


@app.route("/inbox")
@login_required_with_timeout()
@roles_required("admin", "doctor", "patient")
def inbox():
    conn = get_db_connection()
    cursor = conn.cursor()

    cursor.execute(
        """
        SELECT messages.id, users.first_name, users.surname, messages.message, messages.sent_at 
        FROM messages 
        INNER JOIN users ON messages.sender_id = users.id 
        WHERE messages.receiver_id = ? 
        ORDER BY messages.sent_at DESC
    """,
        (session["user_id"],),
    )
    raw_messages = cursor.fetchall()
    conn.close()

    private_key = load_private_key()
    messages = []
    for m in raw_messages:
        messages.append(
            {
                "id": m[0],
                "first_name": decrypt_rsa(m[1], private_key),
                "surname": decrypt_rsa(m[2], private_key),
                "message": decrypt_rsa(m[3], private_key),
                "sent_at": m[4],
            }
        )

    return render_template("inbox.html", messages=messages)


@app.route("/outbox")
@login_required_with_timeout()
@roles_required("admin", "doctor", "patient")
def outbox():
    conn = get_db_connection()
    cursor = conn.cursor()

    cursor.execute(
        """
        SELECT messages.id, users.first_name, users.surname, messages.message, messages.sent_at 
        FROM messages 
        INNER JOIN users ON messages.receiver_id = users.id 
        WHERE messages.sender_id = ? 
        ORDER BY messages.sent_at DESC
    """,
        (session["user_id"],),
    )
    raw_messages = cursor.fetchall()
    conn.close()

    private_key = load_private_key()
    messages = []
    for m in raw_messages:
        messages.append(
            {
                "id": m[0],
                "first_name": decrypt_rsa(m[1], private_key),
                "surname": decrypt_rsa(m[2], private_key),
                "message": decrypt_rsa(m[3], private_key),
                "sent_at": m[4],
            }
        )

    return render_template("outbox.html", messages=messages)


@app.route("/send_message", methods=["GET", "POST"])
@login_required_with_timeout()
@roles_required("admin", "doctor", "patient")
def send_message():
    conn = get_db_connection()
    cursor = conn.cursor()

    if request.method == "POST":
        sender_id = session["user_id"]
        receiver_id = request.form["receiver_id"]
        message = request.form["message"]

        public_key = load_public_key()
        encrypted_message = encrypt_rsa(message, public_key)

        # добавляем текущее время отправки
        sent_at = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

        query = "INSERT INTO messages (sender_id, receiver_id, message, sent_at) VALUES (?, ?, ?, ?)"
        cursor.execute(query, (sender_id, receiver_id, encrypted_message, sent_at))
        conn.commit()

        cursor.close()
        conn.close()

        return redirect(url_for("outbox"))

    cursor.execute("SELECT id, first_name, surname FROM users")
    users_raw = cursor.fetchall()
    conn.close()

    private_key = load_private_key()
    users = []
    for u in users_raw:
        users.append(
            {
                "id": u[0],
                "first_name": decrypt_rsa(u[1], private_key),
                "surname": decrypt_rsa(u[2], private_key),
            }
        )

    return render_template("send_message.html", users=users)


@app.route("/upload-excel/<int:patient_id>", methods=["POST"])
@login_required_with_timeout()
@roles_required("admin", "doctor")
def upload_excel(patient_id):
    file = request.files.get("file")
    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)

        patient_folder = os.path.join("server_database/excel_files", str(patient_id))
        os.makedirs(patient_folder, exist_ok=True)

        filepath = os.path.join(patient_folder, filename)
        file.save(filepath)

        flash("Файл Excel успішно завантажено.", "success")
    else:
        flash("Недійсний формат файлу.", "error")

    return redirect(url_for("patient_dashboard", patient_id=patient_id))


@app.route("/edit_excel/<int:patient_id>/<filename>")
@login_required_with_timeout()
@roles_required("admin", "doctor")
def edit_excel(patient_id, filename):
    """
    Открывает страницу редактора и загружает данные из выбранного файла пациента.
    """
    patient_folder = os.path.join("server_database/excel_files", str(patient_id))
    file_path = os.path.join(patient_folder, filename)

    if not os.path.exists(file_path):
        return redirect(url_for("patient_dashboard", patient_id=patient_id))

    try:
        df = pd.read_excel(file_path, engine="openpyxl", header=None)
        data = df.fillna("").values.tolist()
    except Exception as e:
        return redirect(url_for("patient_dashboard", patient_id=patient_id))

    return render_template(
        "edit_excel.html", patient_id=patient_id, filename=filename, table_data=data
    )


@app.route("/edit_excel/<int:patient_id>/<filename>/save", methods=["POST"])
@login_required_with_timeout()
@roles_required("admin", "doctor")
def save_excel(patient_id, filename):
    data = request.get_json().get("table_data")
    if not data:
        return jsonify({"message": "Немає даних для збереження"}), 400

    try:
        file_path = os.path.join(
            "server_database/excel_files", str(patient_id), filename
        )
        df = pd.DataFrame(data)
        df.to_excel(file_path, index=False, header=False, engine="openpyxl")
        return jsonify({"message": "Таблицю збережено успішно."})
    except Exception as e:
        return jsonify({"message": f"Помилка при збереженні: {str(e)}"}), 500


@app.route("/edit_excel/<int:patient_id>/<filename>/download", methods=["POST"])
@login_required_with_timeout()
@roles_required("admin", "doctor")
def download_excel(patient_id, filename):
    data = request.get_json().get("table_data")
    if not data:
        return "Немає даних для експорту", 400

    try:
        df = pd.DataFrame(data)
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine="openpyxl") as writer:
            df.to_excel(writer, index=False, header=False)
        output.seek(0)

        return send_file(
            output,
            as_attachment=True,
            download_name="edited_table.xlsx",
            mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )
    except Exception as e:
        return f"Помилка при створенні файлу: {str(e)}", 500


@login_required_with_timeout()
@roles_required("admin", "doctor", "patient")
@app.route("/download_patient_excel/<int:patient_id>/<filename>")
def download_patient_excel(patient_id, filename):
    folder_path = os.path.join("server_database", "excel_files", str(patient_id))
    safe_filename = secure_filename(filename)

    return send_from_directory(folder_path, safe_filename, as_attachment=True)


@app.route("/patients/<int:patient_id>/create_new_excel", methods=["POST"])
@login_required_with_timeout()
@roles_required("admin", "doctor")
def create_new_excel(patient_id):
    """Создает новый пустой Excel-файл для пациента."""
    data = request.json
    filename = data.get("filename", "new_table.xlsx")

    patient_folder = os.path.join("server_database/excel_files/", str(patient_id))
    os.makedirs(patient_folder, exist_ok=True)

    file_path = os.path.join(patient_folder, filename)

    try:
        df = pd.DataFrame([["", ""], ["", ""]])
        df.to_excel(file_path, index=False, header=False, engine="openpyxl")
        return jsonify({"success": True})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/calendar/<int:patient_id>", methods=["GET"])
@login_required_with_timeout()
@roles_required("admin", "doctor", "patient")
def get_calendar(patient_id):
    conn = get_db_connection()
    cursor = conn.cursor()

    cursor.execute(
        "SELECT id, title, start, end, description FROM calendar_events WHERE patient_id = ?",
        (patient_id,),
    )
    events = cursor.fetchall()
    conn.close()

    formatted_events = [
        {
            "id": event.id,
            "title": event.title,
            "start": event.start.strftime("%Y-%m-%dT%H:%M") if event.start else None,
            "end": event.end.strftime("%Y-%m-%dT%H:%M") if event.end else None,
            "description": event.description,
        }
        for event in events
    ]

    return jsonify(formatted_events)


@app.route("/calendar/<int:patient_id>", methods=["POST"])
@login_required_with_timeout()
@roles_required("admin", "doctor", "patient")
def create_event(patient_id):
    try:
        data = request.get_json()
        title = data.get("title")
        description = data.get("description")
        start_str = data.get("start")
        end_str = data.get("end")

        start = datetime.fromisoformat(start_str)
        end = datetime.fromisoformat(end_str) if end_str else None

        conn = get_db_connection()
        cursor = conn.cursor()

        cursor.execute(
            """
            INSERT INTO calendar_events (patient_id, title, start, end, description)
            VALUES (?, ?, ?, ?, ?)
            """,
            (patient_id, title, start, end, description),
        )
        conn.commit()
        conn.close()

        return jsonify({"success": True}), 201

    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/meddashboard", methods=["GET"])
@login_required_with_timeout()
@roles_required("admin", "doctor")
def meddashboard():
    conn = get_db_connection()
    cursor = conn.cursor()

    search_query = request.args.get("search", "").strip()
    sort_by = request.args.get("sort_by", "first_name")
    sort_order = request.args.get("sort_order", "ASC").upper()

    valid_sort_columns = ["first_name", "surname"]
    if sort_by not in valid_sort_columns:
        sort_by = "first_name"

    if sort_order not in ["ASC", "DESC"]:
        sort_order = "ASC"

    query = "SELECT id, first_name, surname, phone, email FROM users WHERE position = 'patient'"
    params = []

    if search_query:
        flash("Пошук недоступний через шифрування даних", "error")

    query += f" ORDER BY {sort_by} {sort_order}"

    cursor.execute(query, params)
    raw_patients = cursor.fetchall()
    conn.close()

    private_key = load_private_key()
    patients = []
    for row in raw_patients:
        patients.append(
            {
                "id": row[0],
                "first_name": decrypt_rsa(row[1], private_key),
                "surname": decrypt_rsa(row[2], private_key),
                "phone": decrypt_rsa(row[3], private_key),
                "email": decrypt_rsa(row[4], private_key),
            }
        )

    return render_template(
        "meddashboard.html",
        patients=patients,
        search_query=search_query,
        sort_by=sort_by,
        sort_order=sort_order,
    )


@app.route("/patient/<int:patient_id>")
@login_required_with_timeout()
@roles_required("admin", "doctor")
def patient_dashboard(patient_id):
    try:
        conn = get_db_connection()
        cursor = conn.cursor()

        cursor.execute(
            "SELECT first_name, surname, phone, info, photo FROM users WHERE id = ?",
            (patient_id,),
        )
        patient = cursor.fetchone()
        conn.close()

        if not patient:
            flash("Пацієнта не знайдено.", "error")
            return redirect(url_for("meddashboard"))

        private_key = load_private_key()
        decrypted_patient = {
            "first_name": decrypt_rsa(patient[0], private_key),
            "surname": decrypt_rsa(patient[1], private_key),
            "phone": decrypt_rsa(patient[2], private_key),
            "info": decrypt_rsa(patient[3], private_key) if patient[3] else None,
            "photo": patient[4],
        }

        patient_folder = os.path.join("server_database/excel_files", str(patient_id))
        files = os.listdir(patient_folder) if os.path.exists(patient_folder) else []

        return render_template(
            "patient_dashboard.html",
            patient=decrypted_patient,
            patient_id=patient_id,
            files=files,
        )

    except Exception as e:
        flash(f"Помилка бази даних: {e}", "error")
        return redirect(url_for("dashboard"))


@app.route("/upload-document/<int:patient_id>", methods=["POST"])
@login_required_with_timeout()
@roles_required("admin", "doctor")
def upload_document(patient_id):
    file = request.files["file"]
    if file and patient_id:
        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config["UPLOAD_FOLDER"], filename)
        file.save(filepath)

        if filename.endswith(".pdf"):
            text = extract_text_from_pdf(filepath)
        elif filename.endswith(".docx"):
            text = extract_text_from_docx(filepath)
        else:
            flash("Непідтримуваний формат файлу.", "error")
            return redirect(url_for("patient_dashboard", patient_id=patient_id))

        try:
            public_key = load_public_key()
            encrypted_text = encrypt_rsa(text, public_key)

            conn = get_db_connection()
            cursor = conn.cursor()
            cursor.execute(
                "UPDATE users SET info = ? WHERE id = ?",
                (encrypted_text, patient_id),
            )
            conn.commit()
        except Exception as e:
            flash(f"Помилка бази даних: {str(e)}", "error")
        finally:
            conn.close()

        flash("Документ завантажено та інформацію успішно оновлено.", "success")
        return redirect(url_for("patient_dashboard", patient_id=patient_id))

    flash("Файл не вибрано або ID пацієнта відсутній.", "error")
    return redirect(url_for("dashboard"))


@app.route("/run-tkinter/<patient_id>", methods=["POST"])
@login_required_with_timeout()
@roles_required("admin", "doctor", "patient")
def run_tkinter(patient_id):
    patient_folder = os.path.join("server_database/excel_files/", str(patient_id))

    if not os.path.exists(patient_folder):
        return "Помилка: папку пацієнта не знайдено!", 400

    subprocess.Popen(
        [sys.executable, "graph.py", str(patient_id)], start_new_session=True
    )
    return "График!", 200


if __name__ == "__main__":
    app.run(debug=True)
